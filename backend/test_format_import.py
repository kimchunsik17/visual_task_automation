"""서식 파일 → FormatSpec 역변환(format_import, /api/formats/import)의 추출·개명·API 검사."""

from __future__ import annotations

import pathlib
import subprocess
import sys

import pytest

from documents import format_import, hwpx
from documents.format_spec import FormatSpecError

BACKEND_DIR = pathlib.Path(__file__).resolve().parent


# ── hwpx 추출 ─────────────────────────────────────────────────────────────

def _sample_hwpx(tmp_path) -> str:
    path = str(tmp_path / "주간보고 양식.hwpx")
    hwpx.build({
        "title": "주간 업무 보고",
        "blocks": [
            {"type": "paragraph", "text": "작성자: {{작성자}} / 기간: {{period}}"},
            {"type": "table", "columns": ["항목", "내용"],
             "rows": [["부서", "{{작성자}}"], ["요약", "{{summary}}"]]},
            {"type": "page_break"},
            {"type": "paragraph", "text": "끝."},
        ],
    }, path)
    return path


def test_hwpx_roundtrip_extracts_blocks_and_renames_korean_placeholders(tmp_path):
    spec, info = format_import.spec_from_file(_sample_hwpx(tmp_path))

    types = [b["type"] for b in spec["blocks"]]
    assert "table" in types and "page_break" in types
    texts = [b.get("text", "") for b in spec["blocks"] if b["type"] == "paragraph"]
    assert any("주간 업무 보고" in t for t in texts)   # 빌더의 title 문단도 골격으로 온다

    by_label = {f["label"]: f for f in spec["fields"]}
    # 한글 이름은 ASCII 로 바뀌고 원문이 라벨로 남는다. ASCII 이름은 그대로다.
    assert by_label["작성자"]["name"].startswith("field")
    assert by_label["period"]["name"] == "period"
    assert by_label["summary"]["name"] == "summary"

    # 골격의 참조도 함께 바뀌어 있어야 한다 — 표 셀 두 곳의 {{작성자}} 가 같은 새 이름을 쓴다.
    renamed = by_label["작성자"]["name"]
    table = next(b for b in spec["blocks"] if b["type"] == "table")
    assert table["rows"][0][1] == "{{" + renamed + "}}"
    joined = " ".join(texts)
    assert "{{작성자}}" not in joined and "{{" + renamed + "}}" in joined

    assert spec["output"]["default"] == "hwpx"
    assert info["placeholders"] == ["작성자", "period", "summary"]
    assert info["tables"] == 1


def test_hwpx_import_renders_back(tmp_path, monkeypatch):
    """가져온 초안이 그대로 formatNode 렌더까지 통한다 — 역변환의 완결 조건."""
    monkeypatch.chdir(tmp_path)
    spec, _info = format_import.spec_from_file(_sample_hwpx(tmp_path))
    from documents.format_renderer import render_format
    values = {f["name"]: "값" for f in spec["fields"]}
    result = render_format(spec, values, "docx", str(tmp_path / "out.docx"))
    assert (tmp_path / "out.docx").stat().st_size > 0
    assert result["layout"] == "document"


# ── docx 추출 ─────────────────────────────────────────────────────────────

def _sample_docx(tmp_path) -> str:
    from docx import Document
    doc = Document()
    doc.add_heading("회의록", level=1)
    doc.add_paragraph("회의 일시: {{meetingAt}}")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "안건"
    table.cell(0, 1).text = "결정 사항"
    table.cell(1, 0).text = "{{agenda}}"
    table.cell(1, 1).text = "{{decision}}"
    path = str(tmp_path / "회의록.docx")
    doc.save(path)
    return path


def test_docx_extracts_heading_paragraph_table(tmp_path):
    spec, info = format_import.spec_from_file(_sample_docx(tmp_path))

    heading = next(b for b in spec["blocks"] if b["type"] == "heading")
    assert heading["level"] == 1 and heading["text"] == "회의록"
    table = next(b for b in spec["blocks"] if b["type"] == "table")
    assert table["columns"] == ["안건", "결정 사항"]
    assert table["rows"] == [["{{agenda}}", "{{decision}}"]]
    assert {f["name"] for f in spec["fields"]} == {"meetingAt", "agenda", "decision"}
    assert spec["output"]["default"] == "docx"
    assert info["blocks"] >= 3


# ── 거부 경로 ─────────────────────────────────────────────────────────────

def test_rejects_unknown_extension_and_broken_file(tmp_path):
    bad = tmp_path / "양식.txt"
    bad.write_text("그냥 텍스트", encoding="utf-8")
    with pytest.raises(FormatSpecError):
        format_import.spec_from_file(str(bad))

    fake = tmp_path / "가짜.hwpx"
    fake.write_bytes(b"not a zip at all")
    with pytest.raises(FormatSpecError) as exc:
        format_import.spec_from_file(str(fake))
    assert "열지 못했습니다" in str(exc.value)


def test_rejects_empty_document(tmp_path):
    from docx import Document
    doc = Document()
    path = str(tmp_path / "빈문서.docx")
    doc.save(path)
    with pytest.raises(FormatSpecError) as exc:
        format_import.spec_from_file(path)
    assert "찾지 못했습니다" in str(exc.value)


# ── AI 다듬기 (LLM 은 모의 — test_format_node 의 관례) ───────────────────

def test_refine_imported_spec_keeps_name_and_output(tmp_path, monkeypatch):
    import format_studio

    class _FakeStructured:
        def __init__(self, result): self._result = result
        def invoke(self, _messages): return self._result

    class _FakeLLM:
        def __init__(self, result): self._result = result
        def with_structured_output(self, *_a, **_k): return _FakeStructured(self._result)

    refined = format_studio.GeneratedFormatSpec(
        name="AI가 지은 이름", layout="document",
        fields=[format_studio.GeneratedField(name="author", label="작성자", kind="text", required=True)],
        blocks=[format_studio.GeneratedBlock(type="paragraph", text="작성자: {{author}}")])
    monkeypatch.setattr("meta_agent.get_llm", lambda **_k: _FakeLLM(refined))

    draft = {"name": "주간보고 양식", "layout": "document",
             "output": {"default": "hwpx", "allowed": ["hwpx", "docx", "pdf"]},
             "fields": [], "blocks": [{"type": "paragraph", "text": "작성자: 홍길동"}]}
    result = format_studio.refine_imported_spec(draft)
    # 이름·출력 형식은 파일에서 온 초안이 정본이다.
    assert result["name"] == "주간보고 양식"
    assert result["output"] == draft["output"]
    assert result["fields"][0]["name"] == "author"


def test_refine_skips_oversized_documents():
    import format_studio
    draft = {"name": "큰 문서", "fields": [],
             "blocks": [{"type": "paragraph", "text": "가" * 200} for _ in range(120)]}
    with pytest.raises(FormatSpecError) as exc:
        format_studio.refine_imported_spec(draft)
    assert "건너뜁니다" in str(exc.value)


# ── API 통합 (sqlite 서브프로세스 — 저장소 관례) ─────────────────────────

SCENARIO = r'''
import io, os, sys
os.environ["DATABASE_URL"] = sys.argv[1]
sys.path.insert(0, sys.argv[2])
os.chdir(os.path.dirname(sys.argv[1].replace("sqlite:///", "")))

from fastapi.testclient import TestClient
import main, models
from database import SessionLocal

db = SessionLocal()
user = models.User(id=1, google_id="imp-user", email="imp@example.com", name="imp")
db.add(user)
db.commit()
main.app.dependency_overrides[main.get_current_user_required] = lambda: user
main.app.dependency_overrides[main.get_current_user] = lambda: user
client = TestClient(main.app)

from docx import Document
doc = Document()
doc.add_heading("공문 양식", level=1)
doc.add_paragraph("수신: {{receiver}}")
doc.add_paragraph("발신: {{보내는곳}}")
buffer = io.BytesIO()
doc.save(buffer)

# 1) 가져오기 (AI 끔) — 초안이 오고, 한글 자리표시자는 개명돼 있다
res = client.post("/api/formats/import",
                  files={"file": ("공문 양식.docx", buffer.getvalue(),
                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                  data={"use_ai": "0"})
assert res.status_code == 200, res.text
body = res.json()
assert body["ai"] == "off" and body["spec"]["layout"] == "document"
labels = {f["label"]: f["name"] for f in body["spec"]["fields"]}
assert labels["receiver"] == "receiver" and labels["보내는곳"].startswith("field")
assert body["source"]["placeholders"] == ["receiver", "보내는곳"]

# 2) 가져온 초안이 그대로 라이브러리에 저장된다 (스튜디오의 저장 경로)
saved = client.post("/api/formats", json={"name": body["spec"]["name"], "spec": body["spec"]})
assert saved.status_code == 200, saved.text

# 3) 지원 밖 확장자는 422
bad = client.post("/api/formats/import", files={"file": ("x.txt", b"hello", "text/plain")})
assert bad.status_code == 422, bad.text

# 4) 빈 파일은 422
empty = client.post("/api/formats/import", files={"file": ("y.docx", b"", "application/octet-stream")})
assert empty.status_code == 422, empty.text

print("FORMAT IMPORT ALL OK")
'''


def test_format_import_api_end_to_end(tmp_path):
    pytest.importorskip("httpx", reason="FastAPI TestClient requires httpx")
    scenario_path = tmp_path / "format_import_scenario.py"
    scenario_path.write_text(SCENARIO, encoding="utf-8")
    database_url = f"sqlite:///{tmp_path / 'import.db'}"

    result = subprocess.run(
        [sys.executable, str(scenario_path), database_url, str(BACKEND_DIR)],
        cwd=BACKEND_DIR, capture_output=True, text=True, timeout=300,
    )
    assert result.returncode == 0, f"stdout:\n{result.stdout}\n\nstderr:\n{result.stderr[-3000:]}"
    assert "FORMAT IMPORT ALL OK" in result.stdout
