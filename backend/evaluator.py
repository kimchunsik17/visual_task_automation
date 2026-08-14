import asyncio
import base64
import json
import os
import re
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from dotenv import load_dotenv

import models
from graph import run_workflow
from meta_agent import PLACEHOLDER_URL

# posterGeneratorNode/fileModifierNode처럼 실제 결과가 사람이 읽는 텍스트가 아니라 디스크에
# 저장된 파일(이미지/문서)인 노드들 — discordNode의 파일 첨부 감지(AppViewerPage.jsx의
# FILE_PATH_REGEX)와 동일한 패턴으로 실행 결과 문자열 안의 파일 경로를 찾는다. 이 노드들로
# 끝나는 워크플로우는 예전엔 이 경로 문자열 자체를 텍스트 판정관에게 그대로 보여줘서(판정관
# 입장에선 의미 없는 문자열로만 보이니) 항상 낮은 점수를 받았다 — "이미지 생성 노드는 평가할
# 방법이 없다"는 문제의 원인. 이제 파일 종류에 따라 실제로 내용을 확인해서 채점한다.
_UPLOADS_FILE_RE = re.compile(r'uploads/[^\s"\'<>]+')
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
_DOCX_EXTS = {".docx"}
_HWPX_EXTS = {".hwpx"}
_PDF_EXTS = {".pdf"}
FILE_OUTPUT_NODE_TYPES = {"posterGeneratorNode", "fileModifierNode"}

# 카카오톡/디스코드/텔레그램/이메일/구글시트/구글캘린더/Notion처럼 "외부 서비스로 실제 전송·기록"이
# 목적인 노드들. 평가 파이프라인(run_evaluation_pipeline)은 project_id 없이 run_workflow를 돌리기
# 때문에(golden dataset 생성 단계에서는 아직 저장된 프로젝트가 없어 API 센터 키를 조회할 방법이
# 자체가 없고, 이미 저장된 프로젝트를 "평가" 버튼으로 돌릴 때도 실제 메시지 발송/실제 캘린더 등록
# 같은 부작용을 평가 때마다 일으키지 않기 위해 일부러 실제 키로 치환하지 않는다) 이 노드들은 평가
# 환경에서 항상 자격증명 실패로 끝난다. 각 노드 코드가 실패해도 "무엇을 보내려/기록하려 했는지"는
# "[⚠️ ... 실패/오류/...설정되지 않아 ...]" 형태의 경고문으로 남기고 실제 내용은 버리지 않으므로
# (integration_nodes.py 참고), 이 경고문만 분리해서 판정관에게 "이건 평가 환경의 구조적 한계지
# 워크플로우 결함이 아니다"라고 알려주고 남은 실제 내용을 기준으로 채점하게 한다.
ACTION_SEND_NODE_TYPES = {
    "kakaoNode", "discordNode", "telegramNode", "emailNode",
    "googleSheetsNode", "googleCalendarNode", "notionNode",
}
_ACTION_FAILURE_RE = re.compile(r'\[⚠️[^\]]*\]')


def _strip_action_failure_note(text: str):
    """카카오/디스코드/텔레그램/이메일/구글시트/캘린더/Notion 등이 자격증명 실패 시 덧붙이는
    "[⚠️ ... 실패/오류/...설정되지 않아 ...]" 경고문을 분리한다. (본문, 경고문) 튜플을 반환하며,
    경고문이 없으면 (text, "")."""
    if not text:
        return text, ""
    m = _ACTION_FAILURE_RE.search(text)
    if not m:
        return text, ""
    warning = m.group(0)
    body = (text[:m.start()] + text[m.end():]).strip()
    return body, warning


def _find_output_file(text: str) -> Optional[str]:
    if not text:
        return None
    m = _UPLOADS_FILE_RE.search(text)
    if not m:
        return None
    path = m.group(0)
    return path if os.path.exists(path) else None


def _extract_docx_text(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)[:5000]
    except Exception as e:
        return f"(텍스트 추출 실패: {e})"


def _extract_hwpx_text(path: str) -> str:
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        full_text = ""
        with zipfile.ZipFile(path, "r") as zf:
            sec_files = [n for n in zf.namelist() if n.startswith("Contents/section") and n.endswith(".xml")]
            for sec in sorted(sec_files):
                root = ET.fromstring(zf.read(sec))
                for elem in root.iter():
                    if elem.tag.endswith("}t") or elem.tag.endswith(":t"):
                        if elem.text:
                            full_text += elem.text + " "
        return full_text[:5000]
    except Exception as e:
        return f"(텍스트 추출 실패: {e})"


def _extract_pdf_text(path: str) -> str:
    try:
        import fitz
        doc = fitz.open(path)
        text = "".join(page.get_text() for page in doc)
        return text[:5000]
    except Exception as e:
        return f"(텍스트 추출 실패: {e})"

load_dotenv()
import os
has_langfuse = bool(os.getenv('LANGFUSE_PUBLIC_KEY')) and bool(os.getenv('LANGFUSE_SECRET_KEY'))
if has_langfuse:
    from langfuse.langchain import CallbackHandler
    
# --- Pydantic Models for Structured Output ---

class TestCase(BaseModel):
    input: str = Field(description="The input value to be passed to the workflow (simulating dynamicInputNode).")
    expected_behavior: str = Field(description="A clear description of what the workflow should output or do for this input.")
    evaluation_criteria: str = Field(description="Specific aspects to check in the actual output (e.g. 'Must contain a summary', 'Must output in JSON format').")

class DatasetGenerationResult(BaseModel):
    test_cases: List[TestCase] = Field(description="A list of generated test cases based on the workflow's intent (the exact count is specified in the prompt).")

class JudgeScore(BaseModel):
    correctness: int = Field(description="Score 1-10: Does the output meet the expected behavior?")
    completeness: int = Field(description="Score 1-10: Are all required steps and details included?")
    consistency: int = Field(description="Score 1-10: Is the format, tone, and structure appropriate and consistent?")
    error_handling: int = Field(description="Score 1-10: Did it execute without errors or handle invalid inputs well?")
    usefulness: int = Field(description="Score 1-10: Is the final output practically useful for the user?")
    feedback: str = Field(description="Detailed explanation of the scores and what went wrong/right.")

class TestCaseResult(BaseModel):
    test_case: TestCase
    actual_output: str
    error: Optional[str] = None
    scores: Optional[JudgeScore] = None
    total_score: int = 0

class EvaluationReport(BaseModel):
    score: int = Field(description="Overall workflow quality score (0-100).")
    summary: str = Field(description="A brief summary of the workflow's performance across all test cases.")
    suggestions: List[str] = Field(description="Actionable suggestions to improve the workflow design or prompts.")


# --- Evaluator Agents ---

def get_eval_llm(project_id=None, langfuse_handler=None):
    from llm.providers import create_chat_model

    llm = create_chat_model(
        profile="evaluation",
        temperature=0.2,
        required_capabilities={"structured_output"},
    )
    if has_langfuse and langfuse_handler:
        tags = ["evaluation"]
        metadata = {}
        if project_id:
            metadata["langfuse_session_id"] = f"project-{project_id}"
        llm = llm.with_config(callbacks=[langfuse_handler], metadata=metadata, tags=tags)
    return llm

def generate_golden_dataset(title: str, description: str, nodes: list, edges: list, project_id=None, langfuse_handler=None, num_test_cases: int = 3) -> List[TestCase]:
    """Generates test cases based on the workflow's structure and description.
    num_test_cases: 골든 데이터셋 크기. 기본 3(일반 평가용) — 정밀 모드 품질 게이트처럼 생성마다
    반복 호출되는 경로에서는 속도를 위해 1로 줄여서 부른다(호출부 참고)."""
    llm = get_eval_llm(project_id, langfuse_handler).with_structured_output(DatasetGenerationResult)

    # Analyze if there's a dynamic input node to know what to provide
    input_labels = []
    nodes_summary = []
    for node in nodes:
        if node['type'] == 'dynamicInputNode':
            input_labels.append(node.get('data', {}).get('inputLabel', 'Unknown Input'))

        summary = {"type": node.get("type")}
        if "data" in node:
            for key in ["systemPrompt", "inputLabel", "apiEndpoint", "instruction", "url"]:
                if key in node["data"] and node["data"][key]:
                    summary[key] = node["data"][key]
            # url이 PLACEHOLDER_URL 그대로면(실제 URL을 모를 때 채워지는 값 — meta_agent.py 참고)
            # 이미 위에서 truthy라서 summary에 들어가긴 하지만, 그냥 문자열로만 보이면 골든 데이터셋
            # LLM이 "진짜 URL을 호출해서 결과를 받아온다"고 오해하기 쉽다. 명시적으로 표시한다.
            if summary.get("url") == PLACEHOLDER_URL:
                summary["url"] = "(플레이스홀더 — 실제 URL이 채워지지 않아 실행 시 호출을 시도하지 않고 안내 메시지로 대체됨)"
            # 파일 경로/발송 자격증명은 생성 시점에 일부러 빈 문자열로 남겨두는 게 정상이다
            # (사용자가 에디터에서 직접 채워야 함 — NODE_CATALOG 참고). 그런데 이 필드들이 위 목록에
            # 없어서 골든 테스트 생성 LLM이 "비어있다"는 걸 몰랐고, 그 결과 실제로는 절대 나올 수 없는
            # ("PDF를 요약해서 출력한다", "디스코드로 답장이 온다" 같은) 기대값을 만들어 채점했다 —
            # 워크플로우 품질과 무관하게 파일/발송형 워크플로우가 구조적으로 낮은 점수를 받던 원인.
            # 비어있으면 그 사실 자체를 명시해서, 기대값이 "실행 시 값 채우기를 요청해야 한다" 쪽으로
            # 현실적으로 만들어지게 한다.
            placeholder_keys_by_type = {
                "valueNode": ["file_path", "value"],
                "discordNode": ["botToken", "channelId"],
                "kakaoNode": ["accessToken"],
            }
            for key in placeholder_keys_by_type.get(node.get("type"), []):
                if key in node["data"]:
                    val = node["data"][key]
                    summary[key] = val if val else "(비어있음 — 실행 시점에는 값이 없어 실제 동작을 못 함)"
            # posterGeneratorNode/fileModifierNode처럼 결과가 텍스트가 아니라 파일(이미지/문서)인
            # 노드는 골든 데이터셋 LLM이 "이런 문장이 출력돼야 한다"는 식으로 텍스트 기준을 만들면
            # 실행 결과(파일 경로 문자열)와 무조건 안 맞아서 매번 낮은 점수로 채점됐다 — 이 노드가
            # 흐름의 끝(또는 그 근처)이면 기준 자체를 파일의 시각적/내용적 결과로 잡게 미리 알려준다.
            if node.get("type") in FILE_OUTPUT_NODE_TYPES:
                summary["_output_note"] = (
                    "This node's real output is a FILE saved to disk (image or document), not readable "
                    "text — expected_behavior/evaluation_criteria for a flow ending here should describe "
                    "what the file's visual design or written content should contain, NOT expect a plain "
                    "text chat response."
                )
            # 카카오톡/디스코드/텔레그램/이메일/구글시트/캘린더/Notion 등은 이 평가 환경에서 API 자격
            # 증명이 연결되지 않아 실제 전송·기록은 항상 실패한다(ACTION_SEND_NODE_TYPES 정의 참고).
            # 골든 데이터셋 LLM이 "메시지가 실제로 도착한다"/"시트에 실제로 기록된다" 같은 기대값을
            # 만들면 무조건 못 맞으니, 기준을 "전송/기록하려던 내용의 품질"로 잡게 미리 알려준다.
            if node.get("type") in ACTION_SEND_NODE_TYPES:
                summary["_output_note"] = (
                    "This node sends/writes to an external service (KakaoTalk, Discord, Telegram, Email, "
                    "Google Sheets, Google Calendar, or Notion). In THIS evaluation environment, credentials "
                    "are never connected, so the actual send/write will always fail — that failure is a known "
                    "limitation of the evaluation environment, NOT something the workflow should be judged on. "
                    "expected_behavior/evaluation_criteria for a flow ending here should describe the quality of "
                    "the CONTENT that would have been sent/written, NOT whether it actually arrived."
                )
        nodes_summary.append(summary)

    nodes_info = json.dumps(nodes_summary, ensure_ascii=False)
    has_dynamic_input = bool(input_labels)

    if has_dynamic_input:
        input_context = f"This workflow accepts dynamic inputs: {', '.join(input_labels)}."
        critical_rule = (
            "CRITICAL RULE: DO NOT invent business logic, error messages, or constraints that are not "
            "explicitly defined in the workflow description or Nodes Info. If the workflow is a simple LLM "
            "prompt without strict error handling, the 'expected_behavior' for edge cases should reflect the "
            "natural LLM response (e.g., translating numbers just outputs the numbers), NOT a custom error message."
        )
    else:
        # dynamicInputNode가 없으면 default_input(테스트 입력값)을 실제로 소비하는 노드가 하나도 없다
        # (default_input을 읽는 건 dynamicInputNode의 컴파일 코드뿐 — graph.py/node_generators 참고).
        # 즉 실행할 때마다 노드에 이미 박혀 있는 고정 프롬프트/값만 그대로 쓰인다. 그런데 노드의
        # userPrompt 텍스트 자체가 "문장을 입력해 주세요"처럼 입력을 요구하는 문구를 담고 있으면,
        # LLM이 이 지시를 무시하고 "안녕하세요, 어떻게 지내세요?" 같은 구체적인 입력 시나리오를
        # 지어내서 테스트케이스를 만드는 경우가 실제로 관측됐다 — 그렇게 지어낸 입력은 워크플로우에
        # 전달될 방법이 없으니 평가가 절대 통과할 수 없는 기준이 되어버린다. 프롬프트 지시만으로는
        # 이 hallucination을 안정적으로 막지 못해서, input 필드는 아래에서 코드로 강제 고정한다.
        input_context = (
            "This workflow has NO dynamic input mechanism — no node in the graph reads external input at "
            "execution time. Every run is triggered identically. Even if a node's own text (e.g. a prompt "
            "asking the user to provide a sentence) reads as if it expects input, that text is NOT actually "
            "replaced with anything at runtime — it is sent to the LLM completely as-is, every single time."
        )
        critical_rule = (
            "CRITICAL RULE: This workflow has no dynamic input. DO NOT invent example input sentences, "
            "scenarios, or content that the workflow is supposed to receive or react to — it cannot receive "
            f"anything. All {num_test_cases} test case(s) must use the exact same literal input value 'START' "
            "(this will be enforced regardless of what you output). Base 'expected_behavior' and "
            "'evaluation_criteria' only on what the workflow's fixed nodes will deterministically produce on "
            "their own, not on any hypothetical user-provided content."
        )

    prompt = ChatPromptTemplate.from_messages([
        ("system", f"You are an expert QA Engineer. Your job is to create a 'Golden Dataset' of exactly {num_test_cases} test case(s) to evaluate a no-code workflow. "
                   + ("Provide diverse test cases including edge cases.\n\n" if num_test_cases > 1 else "\n\n")
                   + critical_rule +
                   "\n\nAll text (input, expected_behavior, evaluation_criteria) MUST be written in Korean."),
        ("user", "Workflow Title: {title}\nWorkflow Description: {description}\nNodes Info: {nodes_info}\nEdges: {edges_count}\n\nContext: {input_context}\n\nGenerate the test cases.")
    ])

    chain = prompt | llm
    res = chain.invoke({
        "title": title,
        "description": description,
        "nodes_info": nodes_info,
        "edges_count": len(edges),
        "input_context": input_context
    })

    if not has_dynamic_input:
        # 안전망: 프롬프트 지시를 LLM이 그래도 무시했을 경우를 대비해 input을 코드로 강제 고정한다.
        for tc in res.test_cases:
            tc.input = "START"

    return res.test_cases

def _evaluate_image_with_vision(test_case: TestCase, image_path: str, project_id=None, langfuse_handler=None) -> JudgeScore:
    """이미지 파일은 경로 문자열이 아니라 실제 이미지를 판정관에게 보여주고 채점한다 —
    posterGeneratorNode 같은 이미지 생성 노드를 "평가할 방법이 없어서" 항상 낮은 점수를 받던
    문제의 직접적인 해결책. gpt-4o-mini는 비전(이미지 입력)을 지원한다."""
    llm = get_eval_llm(project_id, langfuse_handler).with_structured_output(JudgeScore)
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    ext = os.path.splitext(image_path)[1].lstrip(".").lower()
    mime = "jpeg" if ext in ("jpg", "jpeg") else ext
    messages = [
        SystemMessage(content=(
            "You are an AI Judge evaluating a workflow's execution output. This workflow's job was to "
            "generate an IMAGE FILE — below you are given the actual generated image itself. Look at it "
            "directly (layout, readability, color, whether the requested content actually appears in it) "
            "and score it across 5 criteria (1-10 each). Do NOT penalize it for 'not being text' — a "
            "well-made image IS the correct output for this kind of node. Provide feedback in Korean."
        )),
        HumanMessage(content=[
            {"type": "text", "text": f"Expected Behavior: {test_case.expected_behavior}\nEvaluation Criteria: {test_case.evaluation_criteria}\n\nHere is the actual generated image:"},
            {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{b64}"}},
        ]),
    ]
    return llm.invoke(messages)


def evaluate_test_case(test_case: TestCase, actual_output: str, error: str, project_id=None, langfuse_handler=None) -> JudgeScore:
    """Evaluates a single test case result and returns a score breakdown."""
    file_path = _find_output_file(actual_output)
    ext = os.path.splitext(file_path)[1].lower() if file_path else None

    if file_path and ext in _IMAGE_EXTS:
        try:
            return _evaluate_image_with_vision(test_case, file_path, project_id, langfuse_handler)
        except Exception as e:
            print(f"[evaluator] 이미지 비전 평가 실패, 텍스트 판정으로 대체: {e}")
            # 아래 텍스트 판정 경로로 그냥 흘러간다(judge_output 계산에서 처리).

    judge_output = actual_output
    if file_path:
        extracted = None
        if ext in _DOCX_EXTS:
            extracted = _extract_docx_text(file_path)
        elif ext in _HWPX_EXTS:
            extracted = _extract_hwpx_text(file_path)
        elif ext in _PDF_EXTS:
            extracted = _extract_pdf_text(file_path)
        elif ext in _IMAGE_EXTS:
            extracted = None  # 비전 평가가 위에서 실패했을 때만 여기 도달 — 내용 미리보기가 불가능한 형식으로 취급
        if extracted:
            judge_output = (
                f"(이 워크플로우는 파일을 생성하는 것이 목적이며, 실제로 '{file_path}' 파일이 생성되었습니다. "
                f"그 파일 안의 실제 텍스트 내용은 다음과 같습니다 — 이 내용을 최종 결과로 평가하세요, "
                f"'파일 경로 문자열' 자체를 결과로 보고 낮은 점수를 주지 마세요:)\n\n{extracted}"
            )
        else:
            judge_output = (
                f"(이 워크플로우는 파일을 생성하는 것이 목적이며, 실제로 '{file_path}' 파일이 생성되었습니다 — "
                "파일 형식 특성상 내용 미리보기는 지원하지 않지만, 파일이 정상적으로 만들어졌다는 것 자체는 "
                "코드로 확인됐습니다. 파일 생성이 요청의 핵심 목적을 달성했다는 점을 감안해서 평가하세요."
            )
    else:
        # 카카오톡/디스코드/텔레그램/이메일/구글시트/캘린더/Notion 등 "발송/기록" 노드는 평가 환경에서
        # 항상 API 자격증명이 연결되지 않아 실제 전송·기록 자체는 실패한다(ACTION_SEND_NODE_TYPES 주석
        # 참고) — 그 자체로 감점하지 말고, 노드가 남겨둔 "실제로 시도했던 내용"을 기준으로 평가하도록
        # 경고문과 본문을 분리해서 판정관에게 맥락을 알려준다.
        body, warning = _strip_action_failure_note(actual_output)
        if warning:
            if body:
                judge_output = (
                    "(이 워크플로우는 외부 서비스(카카오톡/디스코드/텔레그램/이메일/구글시트/구글캘린더/Notion 등)로 "
                    "실제 전송·기록하는 것이 목적입니다. 이 평가 환경에서는 API 자격증명이 연결되어 있지 않아 실제 "
                    f"전송·기록 자체는 실패했습니다({warning}) — 이는 평가 환경의 구조적 한계이지 워크플로우 설계의 "
                    "결함이 아니므로 이 연동 실패 자체로 감점하지 마세요. 대신 실제로 전송·기록하려고 시도했던 "
                    f"내용은 다음과 같습니다 — 이 내용의 품질을 기준으로 평가하세요:)\n\n{body}"
                )
            else:
                judge_output = (
                    "(이 워크플로우는 외부 서비스로 실제 전송·기록하는 것이 목적입니다. 이 평가 환경에서는 API "
                    f"자격증명이 연결되어 있지 않아 실제 전송·기록 자체는 실패했습니다({warning}) — 이는 평가 "
                    "환경의 구조적 한계이지 워크플로우 설계의 결함이 아니니 이 실패만으로 낮은 점수를 주지 "
                    "마세요. 전달하려던 내용 자체가 비어있었는지, 워크플로우 구조상 문제가 있는지는 평소처럼 "
                    "평가해도 됩니다."
                )

    llm = get_eval_llm(project_id, langfuse_handler).with_structured_output(JudgeScore)

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an AI Judge evaluating a workflow's execution output. Score it across 5 criteria (1-10 each) and provide feedback. The feedback MUST be written in Korean."),
        ("user", "Expected Behavior: {expected}\nEvaluation Criteria: {criteria}\n\nActual Output:\n{output}\n\nError (if any):\n{error}\n\nEvaluate now.")
    ])

    chain = prompt | llm
    res = chain.invoke({
        "expected": test_case.expected_behavior,
        "criteria": test_case.evaluation_criteria,
        "output": judge_output or "No output",
        "error": error or "None"
    })
    return res

def summarize_evaluation(test_results: List[TestCaseResult], project_id=None, langfuse_handler=None) -> dict:
    """Creates a final summary and score from all test case results."""
    llm = get_eval_llm(project_id, langfuse_handler).with_structured_output(EvaluationReport)
    
    total_score = sum(r.total_score for r in test_results)
    max_possible = len(test_results) * 50
    normalized_score = int((total_score / max_possible) * 100) if max_possible > 0 else 0
    
    results_dump = []
    for i, r in enumerate(test_results):
        results_dump.append(f"Test {i+1}:\nExpected: {r.test_case.expected_behavior}\nActual: {r.actual_output}\nScore: {r.total_score}/50\nFeedback: {r.scores.feedback if r.scores else 'N/A'}\n")
        
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are the Master Evaluator. Review the test results of a workflow and provide a final summary and actionable suggestions to improve it. DO NOT hallucinate. The overall score is already calculated. The output (summary and suggestions) MUST be written in Korean."),
        ("user", "Overall Score: {score}/100\n\nDetailed Results:\n{results}\n\nProvide the summary and suggestions.")
    ])
    
    chain = prompt | llm
    res = chain.invoke({
        "score": normalized_score,
        "results": "\n".join(results_dump)
    })
    
    # Format the final report as a dict
    report_dict = res.dict() if hasattr(res, 'dict') else res.model_dump()
    report_dict['score'] = normalized_score
    report_dict['test_results'] = [
        {
            "input": r.test_case.input,
            "expected": r.test_case.expected_behavior,
            "actual": r.actual_output,
            "error": r.error,
            "score": r.total_score,
            "feedback": r.scores.feedback if r.scores else "Failed to score"
        }
        for r in test_results
    ]
    return report_dict

async def run_evaluation_pipeline(project_id: int, title: str, description: str, nodes: list, edges: list, db, yield_status=None, user_id: int = None, num_test_cases: int = 3):
    """
    Runs the entire evaluation pipeline:
    1. Generate N golden test cases (기본 3, 정밀 모드 품질 게이트처럼 반복 호출되는 경로는
       num_test_cases=1로 줄여서 부른다 — 호출부 참고)
    2. Run the actual workflow with the test case inputs (테스트케이스끼리 서로 독립적이라 병렬 실행)
    3. Evaluate each output using the Judge LLM
    4. Track token usage and save to DB
    """
    # OpenAI 토큰 추적을 위한 컨텍스트 매니저 사용
    try:
        from langchain_community.callbacks import get_openai_callback
        use_cb = True
    except ImportError:
        use_cb = False

    if yield_status: yield_status("생성 중...")

    eval_token_usage = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}

    handler = None
    if has_langfuse:
        handler = CallbackHandler()

    async def _run_and_score(i: int, tc: TestCase) -> TestCaseResult:
        # run_workflow는 DB 세션을 실제로 쓰므로(예: databaseNode), 병렬 실행 시 하나의 세션을
        # 여러 스레드가 동시에 건드리면 안전하지 않다 — 테스트케이스마다 독립된 세션을 새로 연다.
        def _work():
            from database import SessionLocal
            task_db = SessionLocal()
            try:
                inputs = {"default_input": tc.input}
                try:
                    result_text, tokens, logs = run_workflow(nodes, edges, db=task_db, **inputs)
                    error_msg = None
                    if "► Flow 1 Error:" in result_text or "Dynamic Execution Error:" in result_text or "Execution failed:" in result_text:
                        error_msg = result_text
                except Exception as e:
                    result_text = ""
                    error_msg = str(e)

                try:
                    score = evaluate_test_case(tc, result_text, error_msg, project_id, handler)
                    total = sum([score.correctness, score.completeness, score.consistency, score.error_handling, score.usefulness])
                except Exception as e:
                    print(f"Evaluation failed: {e}")
                    score = None
                    total = 0

                return TestCaseResult(test_case=tc, actual_output=result_text, error=error_msg, scores=score, total_score=total)
            finally:
                task_db.close()

        if yield_status:
            yield_status(f"테스트 {i+1} 실행/평가 중...")
        return await asyncio.to_thread(_work)

    async def _run_pipeline():
        try:
            if yield_status:
                yield_status("테스트 케이스 생성 중...")
            # generate_golden_dataset/summarize_evaluation은 내부적으로 LLM을 동기(blocking)로
            # 호출한다(chain.invoke) — uvicorn이 --workers 1(단일 워커)로 떠 있어서, 이걸 await 없이
            # 그냥 부르면 그 호출이 끝날 때까지 이벤트 루프 전체가 멈춰서 다른 모든 사용자의 요청까지
            # 같이 멈춘다("서버가 멈춘 것 같다"는 증상의 실제 원인 — 정밀 모드 품질 게이트가 워크플로우
            # 생성마다 이 경로를 타므로 흔하게 재발할 수 있었다). _run_and_score의 실제 실행 부분은
            # 이미 asyncio.to_thread로 스레드에 넘기고 있었는데, 데이터셋 생성/요약 두 곳만 빠져 있었다.
            test_cases = await asyncio.to_thread(
                generate_golden_dataset, title, description, nodes, edges, project_id, handler, num_test_cases
            )
        except Exception as e:
            print(f"Failed to generate dataset: {e}")
            return {"error": "Failed to generate dataset"}

        results = await asyncio.gather(*[_run_and_score(i, tc) for i, tc in enumerate(test_cases)])

        if yield_status:
            yield_status("최종 리포트 생성 중...")
        return await asyncio.to_thread(summarize_evaluation, list(results), project_id, handler)

    # 토큰 추적 래퍼
    if use_cb:
        from langchain_community.callbacks import get_openai_callback
        with get_openai_callback() as cb:
            report = await _run_pipeline()
        eval_token_usage = {
            "input_tokens": cb.prompt_tokens,
            "output_tokens": cb.completion_tokens,
            "total_tokens": cb.total_tokens,
        }
    else:
        report = await _run_pipeline()

    if isinstance(report, dict) and "error" not in report:
        # 평가 토큰을 FlowExecutionLog에 저장
        try:
            import json as _json
            db_log = models.FlowExecutionLog(
                user_id=user_id,
                project_id=project_id,
                payload="Evaluation Pipeline",
                result=f"Score: {report.get('score', 0)}/100",
                total_tokens=eval_token_usage["total_tokens"],
                token_usage_details=eval_token_usage,
                status="evaluation",
            )
            db.add(db_log)
            db.commit()
        except Exception as e:
            print(f"Failed to save evaluation token log: {e}")
            db.rollback()

        # token_usage를 report에 포함
        report["token_usage"] = eval_token_usage

    if handler and hasattr(handler, 'flush'):
        handler.flush()

    return report


async def run_evaluation_with_autofix(
    project_id: int, title: str, description: str, nodes: list, edges: list, db,
    yield_status=None, user_id: int = None, threshold: int = 70, max_attempts: int = 3,
) -> dict:
    """평가 점수가 threshold 미만이면, 개선 제안을 메타 에이전트(run_agent_turn)에 다시 넣어
    워크플로우를 자동 수정하고 재평가한다. threshold 통과 또는 max_attempts 도달 시 종료.

    반환: run_evaluation_pipeline과 동일한 report dict + attempts(시도별 점수 이력) +
    graph_data(최종 반영해야 할 노드/엣지, 수정이 일어났을 때만 원본과 달라짐).
    """
    from meta_agent import run_agent_turn

    current_nodes, current_edges = nodes, edges
    attempts = []
    report = None
    total_autofix_tokens = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}

    for attempt in range(1, max_attempts + 1):
        if yield_status:
            yield_status(f"평가 중... (시도 {attempt}/{max_attempts})")
        report = await run_evaluation_pipeline(
            project_id, title, description, current_nodes, current_edges, db, yield_status, user_id
        )
        if "error" in report:
            break

        score = report.get("score", 0)
        attempts.append({"attempt": attempt, "score": score})

        if score >= threshold or attempt >= max_attempts:
            break

        if yield_status:
            yield_status(f"기준({threshold}점) 미달({score}점) — 자동 수정 시도 중...")

        suggestions = report.get("suggestions", [])
        fix_message = (
            f"방금 이 워크플로우가 자동 평가에서 {score}/100점을 받아 기준({threshold}점)에 못 미쳤습니다. "
            "아래 개선 제안을 반영해서 워크플로우를 수정해줘:\n" + "\n".join(f"- {s}" for s in suggestions)
        )
        try:
            _, fixed_graph_data, autofix_tokens, _ = await run_agent_turn(
                {"title": title, "description": description, "nodes": current_nodes, "edges": current_edges},
                fix_message,
                thread_id=f"project-{project_id}-autofix",
            )
            current_nodes = fixed_graph_data.get("nodes", current_nodes)
            current_edges = fixed_graph_data.get("edges", current_edges)
            for k in total_autofix_tokens:
                total_autofix_tokens[k] += autofix_tokens.get(k, 0)
        except Exception as e:
            print(f"Autofix attempt {attempt} failed: {e}")
            break

    if report is not None:
        report["attempts"] = attempts
        report["graph_data"] = {"nodes": current_nodes, "edges": current_edges}
        report["autofix_token_usage"] = total_autofix_tokens

    return report
