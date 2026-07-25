"""
template_generator.py — templateAnalyzerNode/fileModifierNode가 참조하는 서식 파일이
실제로 업로드되지 않았을 때, 그 자리에 즉석에서 채울 수 있는 실제 파일을 만들어 넣는다.

배경: 워크플로우 생성 챗봇은 template_path에 "자기소개서_템플릿.hwpx" 같은 그럴싸한 파일명만
지어낼 뿐, 실제 파일을 만들 방법이 없었다. 그래서 사용자가 에디터에서 직접 파일을 업로드하지
않으면 항상 "No such file or directory" 에러가 났다(project 30에서 실제로 겪음). 이 모듈은
그 자리를 "빈칸에 {{key}} 표시가 있는 진짜 파일"로 즉석에서 채워서, 이후 로직(템플릿 스캔,
채워넣기)이 실제로 업로드된 파일과 똑같이 동작하게 만든다.

.hwpx는 python-hwpx 라이브러리로 만든다 — 이 서버는 리눅스라 실제 "한글" 프로그램으로 열어서
검증할 방법은 없지만, 이 라이브러리 자체의 validate_editor_open_safety()가 자체 검증한 결과
blocking 오류 없이 통과하는 것을 확인했다(직접 만든 zip+XML보다 훨씬 신뢰도가 높다).
.docx는 python-docx로, .pdf는 PyMuPDF(fitz)로 직접 렌더링한다 — PDF는 hwpx/docx처럼
"빈칸이 있는 서식 파일을 나중에 채우는" 방식이 아니라, 값이 정해지면 그 자리에서 바로 완성된
문서를 만드는 방식이라(PDF는 텍스트를 안전하게 찾아 바꾸기 어려운 포맷이라 서식-채우기 패턴이
잘 안 맞는다) fileModifierNode에서 output 확장자가 .pdf면 이 경로를 탄다.
"""
import os


def generate_hwpx_template(save_path: str, fields: list, title: str = "문서") -> None:
    from hwpx.document import HwpxDocument

    doc = HwpxDocument.new()
    doc.add_paragraph(title)
    for key in fields:
        doc.add_paragraph(f"{key}: {{{{{key}}}}}")
    os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
    doc.save_to_path(save_path)


def extract_template_keys(path: str) -> set:
    """기존 파일(.hwpx/.docx)에 이미 있는 {{key}} 자리표시자 이름들을 뽑아낸다.

    배경: fileModifierNode가 "템플릿이 없으면 즉석 생성"할 때 os.path.exists()만 보고
    판단하면, 챗봇이 지어낸 파일명이 우연히(또는 이전 실행에서 만들어져) 이미 존재하는
    경우 그 낡은 템플릿을 그대로 재사용해버린다 — 그 안의 {{key}}가 이번에 채우려는 값의
    키와 다르면 하나도 안 채워진 채 남는다(실제로 겪음: 자기소개서 hwpx에 {{introduction}}
    등이 그대로 남음). 그래서 재사용 전에 기존 파일의 실제 필드 이름을 먼저 확인해서, 지금
    채우려는 키들과 얼마나 겹치는지 판단할 수 있게 한다.
    """
    import re

    full_text = ""
    ext = path.lower()
    try:
        if ext.endswith(".hwpx"):
            import zipfile
            import xml.etree.ElementTree as ET

            with zipfile.ZipFile(path, "r") as zf:
                sec_files = [n for n in zf.namelist() if n.startswith("Contents/section") and n.endswith(".xml")]
                for sec in sorted(sec_files):
                    root = ET.fromstring(zf.read(sec))
                    for elem in root.iter():
                        if elem.tag.endswith("}t") or elem.tag.endswith(":t"):
                            if elem.text:
                                full_text += elem.text + " "
        elif ext.endswith(".docx"):
            from docx import Document

            doc = Document(path)
            for p in doc.paragraphs:
                full_text += p.text + " "
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
    except Exception:
        return set()

    return {k.strip() for k in re.findall(r"\{\{([^}]+)\}\}", full_text)}


def generate_docx_template(save_path: str, fields: list, title: str = "문서") -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for key in fields:
        doc.add_paragraph(f"{key}: {{{{{key}}}}}")
    os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
    doc.save(save_path)


def render_pdf_document(save_path: str, data: dict, title: str = "문서") -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    x, y = 72, 72
    line_height = 22
    page.insert_text((x, y), str(title), fontsize=16, fontname="korea")
    y += line_height * 2
    for key, value in data.items():
        text = f"{key}: {value}"
        for line in _wrap_text(text, 60):
            if y > page.rect.height - 72:
                page = doc.new_page()
                y = 72
            page.insert_text((x, y), line, fontsize=11, fontname="korea")
            y += line_height
    os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
    doc.save(save_path)


def _wrap_text(text: str, width: int) -> list:
    lines = []
    current = ""
    for ch in text:
        current += ch
        if len(current) >= width and ch in " ,.":
            lines.append(current)
            current = ""
    if current:
        lines.append(current)
    return lines or [""]
