"""documents/docx_builder.py — DocumentSpec(hwpx builder 와 같은 5블록)을 .docx 로 렌더한다.

hwpx builder 와 계약을 맞춘다: 지원하지 않는 블록은 조용히 빠뜨리지 않고 SpecError 로 실패,
이미지는 artifactId + image_loader 주입으로만.
"""

from __future__ import annotations

import io
import os
from typing import Any, Dict, Optional

from .hwpx.builder import (
    IMAGE_FORMATS, MAX_BLOCKS, MAX_TABLE_CELLS, SUPPORTED_BLOCKS, ImageLoader, SpecError,
    UnsupportedFeature,
)

_HEADING_PT = {1: 18, 2: 15, 3: 13}


def build(spec: Dict[str, Any], output_path: str,
          image_loader: Optional[ImageLoader] = None) -> Dict[str, Any]:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Mm, Pt

    if not isinstance(spec, dict):
        raise SpecError("DocumentSpec 은 객체여야 합니다.")
    blocks = spec.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        raise SpecError("blocks 가 비어 있습니다.")
    if len(blocks) > MAX_BLOCKS:
        raise SpecError(f"블록이 너무 많습니다({len(blocks)}개, 상한 {MAX_BLOCKS}개).")

    document = Document()

    for index, block in enumerate(blocks):
        if not isinstance(block, dict):
            raise SpecError(f"blocks[{index}] 는 객체여야 합니다.")
        btype = block.get("type")
        if btype not in SUPPORTED_BLOCKS:
            raise UnsupportedFeature(
                f"blocks[{index}].type {btype!r} 는 지원하지 않습니다({', '.join(SUPPORTED_BLOCKS)}).")

        if btype == "heading":
            level = block.get("level", 1)
            if level not in _HEADING_PT:
                raise SpecError(f"heading.level 은 {sorted(_HEADING_PT)} 중 하나여야 합니다.")
            text = block.get("text")
            if not isinstance(text, str) or not text:
                raise SpecError("heading.text 가 비어 있습니다.")
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(_HEADING_PT[level])

        elif btype == "paragraph":
            text = block.get("text")
            if text is None:
                text = ""
            if not isinstance(text, str):
                raise SpecError("paragraph.text 는 문자열이어야 합니다.")
            document.add_paragraph(text)

        elif btype == "table":
            columns = block.get("columns")
            rows = block.get("rows")
            if not isinstance(columns, list) or not columns:
                raise SpecError("table.columns 는 비어 있지 않은 배열이어야 합니다.")
            if not isinstance(rows, list):
                raise SpecError("table.rows 는 배열이어야 합니다.")
            width = len(columns)
            total_cells = width * (len(rows) + 1)
            if total_cells > MAX_TABLE_CELLS:
                raise SpecError(f"표가 너무 큽니다({total_cells}칸, 상한 {MAX_TABLE_CELLS}칸).")
            table = document.add_table(rows=len(rows) + 1, cols=width)
            table.style = "Table Grid"
            for column_index, header in enumerate(columns):
                cell = table.cell(0, column_index)
                cell.text = "" if header is None else str(header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for row_index, row in enumerate(rows, start=1):
                if not isinstance(row, list) or len(row) != width:
                    raise SpecError(f"table.rows[{row_index - 1}] 의 칸 수가 columns 와 다릅니다.")
                for column_index, value in enumerate(row):
                    table.cell(row_index, column_index).text = "" if value is None else str(value)

        elif btype == "image":
            artifact_id = block.get("artifactId")
            if not artifact_id:
                raise SpecError("image 블록은 artifactId 로만 지정할 수 있습니다(경로·URL 은 받지 않습니다).")
            if image_loader is None:
                raise SpecError("이 실행에서는 이미지를 넣을 수 없습니다(Artifact 해석기가 없습니다).")
            data, image_format = image_loader(str(artifact_id))
            image_format = str(image_format or "").lower().lstrip(".")
            if image_format not in IMAGE_FORMATS:
                raise SpecError(f"지원하지 않는 이미지 형식입니다: {image_format or '(알 수 없음)'}")
            kwargs = {}
            if block.get("widthMm") is not None:
                try:
                    kwargs["width"] = Mm(float(block["widthMm"]))
                except (TypeError, ValueError):
                    raise SpecError("image.widthMm 는 숫자여야 합니다.") from None
            document.add_picture(io.BytesIO(data), **kwargs)

        else:  # page_break
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    document.save(output_path)
    return {"path": output_path, "blocks": len(blocks)}
